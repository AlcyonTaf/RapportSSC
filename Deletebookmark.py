# -*- coding: utf-8 -*-
import itertools

import docx
from lxml import etree as et
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()

from docx.table import _Cell, Table
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.oxml.shared import qn
from docx.oxml.text.paragraph import CT_P
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def is_root(element):
    if element.getparent() is None:
        return True
    else:
        return False


def delete_bookmark(
        doc,
        bookmark_name,
        header=False):
    doc_element = (
        doc.sections[0].header.part._element
        if header is True
        else doc._part._element
    )
    # print(et.tostring(doc._part._element, pretty_print=True)) # Pour obtenir tout le document en xml
    bookmarks_list = doc_element.findall(".//" + qn("w:bookmarkStart"))
    bookmarks_list_end = doc_element.findall(".//" + qn("w:bookmarkEnd"))

    find = False
    Element_To_Remove = []
    for element in doc_element.iter():
        #print(element.tag)
        if element.tag == qn("w:bookmarkStart") and element.get(qn("w:name")) == 'BK_Delete_1':
            print("%s - %s - %s" % (element.tag, element.text, element.get(qn("w:name"))))
            print(element.get(qn("w:id")))
            Bk_id = element.get(qn("w:id"))
            print(is_root(element.getparent()))
            find = True


        if find:
            elem_loop = element
            print(element.tag + " - " + qn("w:body>"))

            findbody = False
            while True:
                print(elem_loop.tag + " - " + qn("w:body>"))
                if elem_loop.getparent().tag == qn("w:body"):
                    print('trouvé')
                    Element_To_Remove.append(elem_loop)
                    break
                else:
                    elem_loop = elem_loop.getparent()
                    #print(elem_loop)


            if element.tag == qn("w:bookmarkEnd") and element.get(qn("w:id")) == Bk_id:
                print('Fin de notre BK')
                #print(is_root(element.getparent()))
                #print(is_root(element.getparent().getparent()))
                find = False


    #Apparement il ya des doublons, on les vires :
    Element_To_Remove = list(dict.fromkeys(Element_To_Remove))
    print(Element_To_Remove)

    body = doc_element.getchildren()[0]

    for elem_supp in Element_To_Remove:
        print(elem_supp)
        print(elem_supp.getparent().tag)
        body.remove(elem_supp)


doc = Document(".\ExtractionLims\Test Rapport Essai corrosion - modif.docx")

delete_bookmark(doc, 'BK_Delete_1')


doc.save("resultdelete.docx")