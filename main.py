# -*- coding: utf-8 -*-
import itertools

import docx
from lxml import etree as et
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()

from docx.table import _Cell, Table
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.oxml.shared import qn
from docx.oxml.text.paragraph import CT_P
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Le but va etre de mettre en forme le rapport des essais SSC de la corrosion
# TOdo : Premiere chose a faire : Parcourir toutes les tables provisoire pour extraire les données et les stocker, ensuite on les supprimes


def get_data_from_table():
    """
    Le but est de récupérer le contenue des tables provisoires et ensuite les supprimer
    Position des tables :
    :return:
    """
    # Nom et position des tableaux à récupérer
    table_list = {'items_essai': 1, 'item_tth': 2, 'item_parent_tth' : 3, 'items_tole': 4, 'item_tole_TTH': 5,
                  'conditions_essais': 6, 'resultats_essais': 8}
    table_valeur = {}
    # Todo : suppresion des tableaux
    for nom_table, position in table_list.items():
        # Exemple pour recupérer les données d'un tableau
        get_table = doc.tables[position]

        # pour mettre dans une liste :
        data = [['' for i in range(len(get_table.columns))] for j in range(len(get_table.rows))]
        for y, row in enumerate(get_table.rows):
            for j, cell in enumerate(row.cells):
                if cell.text:
                    data[y][j] = cell.text

        table_valeur[nom_table] = data

    # Maintenant que l'on a tous récupérer on supprime les tables du documents :
    increment = 0
    for nom_table, position in table_list.items():
        # print("delete :" + nom_table)
        # print("avant: " + str(position))
        # print("new :" + str(position-increment))
        delete_table = doc.tables[position - increment]
        delete_table._element.getparent().remove(delete_table._element)
        increment += 1

    return table_valeur


def clean_table_data():
    """
    On va faire le trie/nettoyage dans certaines tables ici
    :param dict_tables:
    :return: dict nettoyer
    """
    dict_tables = get_data_from_table()

    # resultats_essais : Il faut trouvé la ref client dans la table items_essai et l'ajouter dans cette table avec de nettoyer items_essais
    resultats_essais = dict_tables['resultats_essais']
    items_essai = dict_tables['items_essai']
    #print(resultats_essais)
    #print(items_essai)
    for row in resultats_essais[1:]:
        #print(row[0])
        for rowx in items_essai[1:]:
            try:
                #print(rowx.index(row[0]))
                row[0] = rowx[1]
            except ValueError:
                pass
            # for index, item in enumerate(rowx):
            #    print(str(index) + " -" + str(item))
        #test = [index for (index, item) in enumerate(items_essai[1:]) if item == row[0]]
        #print(test)
    #print(resultats_essais)



    # items_parent :
    # On va supprimer les doublons puis vérifier que la table contient bien uniquement 2 lignes, sinon c'est qu'il y a plusieurs items parents!
    # Todo : Voir pour vérifier également que cette table contient bien des données.
    # Todo : vérifier les forms des items parents
    # 26-07-22 : Désactivé car cette table n'est plus utile
    # items_parent = dict_tables['items_parent']
    # items_parent = list(items_parent for items_parent, _ in itertools.groupby(items_parent))
    # try:
    #     if len(items_parent) == 2:
    #         dict_tables['items_parent'] = items_parent
    #         # print(dict_tables['items_parent'])
    #     else:
    #         # Plus ou moins de 2 ligne = probleme!!
    #         raise ValueError
    # except ValueError:
    #     print("Les items n'ont pas tous le même parents!")

    # items_essai :
    # On va lui donner la forme du tableau de destination :
    # Normalement les informations seront toujours a la même position, sauf si on change l'export de teexma
    # Concatenation des dimensions => Position 2
    # Concaténation des positons => Position 0
    # Ref client => Position 1
    items_essai = dict_tables['items_essai']
    # print(items_essai)
    temp_items_essais = []
    for row in items_essai[1:]:
        ref_item = row[1]
        dimensions = row[5] + "*" + row[6] + "*" + row[7]
        position = row[2] + " / " + row[3]
        temp_items_essais.append([position, ref_item, dimensions])

    dict_tables['items_essai'] = temp_items_essais

    # print(dict_tables['items_essai'])

    # conditions_essais :
    conditions_essais = dict_tables['conditions_essais']
    # On supprime la 1er colonne de chaque ligne car ne sert a rien et comme différente empeche la suppresion des doublon
    conditions_essais = [row[1:] for row in conditions_essais]
    conditions_essais = list(conditions_essais for conditions_essais, _ in itertools.groupby(conditions_essais))
    try:
        if len(conditions_essais) == 2:
            dict_tables['conditions_essais'] = conditions_essais
            # print(dict_tables['items_parent'])
        else:
            # Plus ou moins de 2 ligne = probleme!!
            raise ValueError
    except ValueError:
        print("Les items n'ont pas tous les mêmes conditions d'essais")

    return dict_tables


# Trouver un bookmark : https://stackoverflow.com/questions/24965042/python-docx-insertion-point
def get_bookmark_par_element(document2, bookmark_name):
    """
    Return the named bookmark parent paragraph element. If no matching
    bookmark is found, the result is '1'. If an error is encountered, '2'
    is returned.
    """
    doc_element = document2.part._element
    bookmarks_list = doc_element.findall('.//' + qn('w:bookmarkStart'))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn('w:name'))
        if name == bookmark_name:
            par = bookmark.getparent()
            if not isinstance(par, docx.oxml.CT_P):
                return 2
            else:
                return par
    return 1


def bookmark_text(
        doc,
        bookmark_name,
        text,
        underline=False,
        italic=False,
        bold=False,
        style=None,
        header=False,
        font_name=None,
        font_size=None):
    doc_element = (
        doc.sections[0].header.part._element
        if header is True
        else doc._part._element
    )
    bookmarks_list = doc_element.findall(".//" + qn("w:bookmarkStart"))
    for bookmark in bookmarks_list:
        name = bookmark.get(qn("w:name"))
        if name == bookmark_name:
            par = bookmark.getparent()

            # => Rajouter pour supprimer par position start et end
            pos_BK_start = par.index(bookmark)
            # print(pos_BK_start)
            BK_id = bookmark.get(qn("w:id"))  # => Rajouter pour supprimer par position start et end
            bookmark_end_list = par.findall(
                ".//" + qn("w:bookmarkEnd"))  # => Rajouter pour supprimer par position start et end
            for BK_end in bookmark_end_list:
                if BK_end.get(qn("w:id")) == BK_id:
                    pos_BK_end = par.index(BK_end)

            if not isinstance(par, CT_P):
                return False
            else:
                # for elem in par.iter():
                #     print("%s - %s" % (elem.tag, elem.text))
                # print(et.tostring(par, pretty_print=True))
                # Todo : Voir si on ne peut pas plutot modifier le text du bookmark
                i = par.index(bookmark)  # - 1 Désactivé car insert a deux  index avant le bk
                # print(et.tostring(par[i], pretty_print=True))
                p = doc.add_paragraph()
                run = p.add_run(text, style)
                run.underline = underline
                run.italic = italic
                run.bold = bold
                run.font.size = font_size
                run.font.name = font_name
                par.insert(i, run._element)
                # print(et.tostring(par, pretty_print=True))
                p = p._element
                p.getparent().remove(p)
                p._p = p._element = None
                # Essai pour suppresion bookmarks + text
                for z in range(pos_BK_end, pos_BK_start, -1):
                    # print(z)
                    asup = par[z]
                    #
                    # print(et.tostring(asup, pretty_print=True))
                    par.remove(asup)
                return True
    return False


def replace_bk_by_value(bk_dict, table_source):
    for bk_name, position in bk_dict.items():
        # cas particulier
        # Nuance peut etre a 2 endroit, on regarde si le 1er est vide si oui on prend le 2eme
        if bk_name == 'BK_Item_Nuance':
            if result[table_source][1][position[0]]:
                add_bk = bookmark_text(doc, bk_name, result[table_source][1][position[0]], font_name='Arial',
                                       font_size=127000)
            elif not result[table_source][1][position[0]] and result[table_source][1][position[1]]:
                add_bk = bookmark_text(doc, bk_name, result[table_source][1][position[1]], font_name='Arial',
                                       font_size=127000)
            else:
                add_bk = bookmark_text(doc, bk_name, "N/A", font_name='Arial', font_size=127000)
        # Vérification que l'item parent de la Tole est bien la coulée
        elif bk_name == 'BK_Item_Coulee':
            if result[table_source][1][7] == 'Coulée':
                add_bk = bookmark_text(doc, bk_name, result[table_source][1][position], font_name='Arial',
                                       font_size=127000)
            else:
                add_bk = bookmark_text(doc, bk_name, "N/A", font_name='Arial', font_size=127000)
        else:
            if not result[table_source][1][position]:
                add_bk = bookmark_text(doc, bk_name, "N/A", font_name='Arial', font_size=127000)
            else:
                add_bk = bookmark_text(doc, bk_name, result[table_source][1][position], font_name='Arial',
                                       font_size=127000)

        try:
            if not add_bk:
                raise ValueError
        except ValueError:
            print("Probleme lors de l'insertion d'un BK!")


def delete_bookmark(
        doc,
        bookmark_name,
        header=False):
    doc_element = (
        doc.sections[0].header.part._element
        if header is True
        else doc._part._element
    )

    find = False
    element_to_remove = []
    for element in doc_element.iter():
        if element.tag == qn("w:bookmarkStart") and element.get(qn("w:name")) == bookmark_name:
            #print("%s - %s - %s" % (element.tag, element.text, element.get(qn("w:name"))))
            #print(element.get(qn("w:id")))
            Bk_id = element.get(qn("w:id"))
            find = True

        if find:
            elem_loop = element
            # print(element.tag + " - " + qn("w:body>"))
            # On considere que tous les éléments a supprimer sont des enfants de w:body
            # On va donc remonter depuis l'element en cours pour trouvé le parent qui est enfant de w:body
            while True:
                # print(elem_loop.tag + " - " + qn("w:body>"))
                if elem_loop.getparent().tag == qn("w:body"):
                    # print('trouvé')
                    element_to_remove.append(elem_loop)
                    break
                else:
                    # On remonte au parent
                    elem_loop = elem_loop.getparent()
                    # print(elem_loop)

            if element.tag == qn("w:bookmarkEnd") and element.get(qn("w:id")) == Bk_id:
                # print('Fin de notre BK')
                # print(is_root(element.getparent()))
                # print(is_root(element.getparent().getparent()))
                find = False

    # On supprime les doublons
    element_to_remove = list(dict.fromkeys(element_to_remove))

    body = doc_element.getchildren()[0]
    # On boucle sur les éléments a supprimer
    for elem_supp in element_to_remove:
        body.remove(elem_supp)

# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    doc = Document(".\ExtractionLims\Test Rapport Essai corrosion - V3.docx")
    all_paras = doc.paragraphs
    # print(len(all_paras))

    # print(get_data_from_table())
    result = clean_table_data()

    # On complete la table Items_Essais
    table_items_essais = doc.tables[0]
    # essai pour récupérer les info de styles de la table
    # for cell in table_items_essais.row_cells(1):
    #     for paragraph in cell.paragraphs:
    #         for run in paragraph.runs:
    #             print(run.font)
    #             font = run.font
    for row_item in result['items_essai']:
        new_row = table_items_essais.add_row().cells
        for i, val in enumerate(row_item):
            new_row[i].text = val
            for paragraph in new_row[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = 127000

    # Todo : Remplacer les bookmarks
    ###### ATTENTION ######
    # Il ne faut pas qu'il y ai des demandent de correction dans le word, il faut tous traiter ou ignorer
    # Remplacement BK de l'item parent :
    # On créer un dict avec le schema nomBK : Position tableau

    BK_item_tole = {'BK_Item_Nuance': [2, 3], 'BK_Item_UM': 1, 'BK_Item_Coulee': 5,
                    'BK_Item_EP_UM': 4}
    replace_bk_by_value(BK_item_tole, table_source='items_tole')

    # Gestion TTH sur Tole :
    # On va concatener les informations du tableau des TTHs produit.
    # Une colonne par type de TTH, 1er ligne = nom, 2eme ligne = température, 3eme ligne = durée
    tth_list = list(itertools.zip_longest(*result['item_tole_TTH']))
    #print(tth_list)
    tth_temp = []
    for row in tth_list:
        tth_temp.append('-'.join(row))

    tth_tole = ' / '.join(tth_temp)
    bookmark_text(doc, 'BK_Item_TTH_Parent', tth_tole, font_name='Arial',
                  font_size=127000)

    # Remplacement BK des conditions de l'essai :
    BK_conditions_essais = {'BK_Condition_Solution': 0, 'BK_Condition_Gaz_Degazage': 1, 'BK_Condition_Gaz_Essai': 2,
                            'BK_Condition_ph': 3, 'BK_Condition_ph_Saturation': 4, 'BK_Condtion_ph_fin': 5,
                            'BK_Condition_Temp': 6, 'BK_Condition_Duree': 7, 'BK_Condition_Limite_Reelle': 8,
                            'BK_Condition_Limite_Garantie': 9, 'BK_Condition_Contrainte': 10,
                            'BK_Condition_Method': 11, 'BK_Condition_Examen': 12}
    replace_bk_by_value(BK_conditions_essais, table_source='conditions_essais')


    # Pour finir on supprime les BK_delete_x
    # Todo : Penser a modifier le templace LIMS pour ajouter les BK_delete_x
    delete_bookmark(doc, 'BK_Delete_1')


    # test = get_bookmark_par_element(doc, "BK_Condition_Solution")
    # print(test.r)

    # test = bookmark_text(doc, "BK_Condition_Solution", " 1111", font_name='Arial', font_size=127000)

    # Exemple pour recupérer les données d'un tableau
    # tb = doc.tables[1]

    # pour mettre dans une liste :
    # list = [['' for i in range(len(tb.columns))] for j in range(len(tb.rows))]
    # for y, row in enumerate(tb.rows):
    #     for j, cell in enumerate(row.cells):
    #         if cell.text:
    #             list[y][j] = cell.text
    #
    # print(list)
    # print(type(list))

    # Pour recuper le contenue dans un dict avec nom de colonne comme clef
    # data =[]
    # for i, row in enumerate(tb.rows):
    #     print(i)
    #     text = (cell.text for cell in row.cells)
    #
    #     # Establish the mapping based on the first row
    #     # headers; these will become the keys of our dictionary
    #     if i == 0:
    #         keys = tuple(text)
    #         print(keys)
    #         continue
    #
    #     # Construct a dictionary for this row, mapping
    #     # keys to values for this row
    #     row_data = dict(zip(keys, text))
    #     print(row_data)
    #     data.append(row_data)
    #
    # print(data)




    doc.save("result.docx")
