import itertools

import docx
from lxml import etree as et
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()

from docx.oxml.shared import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def bookmark_text(
        doc,
        bookmark_name,
        text,
        underline=False,
        italic=False,
        bold=False,
        style=None,
        header=False,
        font_name=None,
        font_size=None):
    doc_element = (
        doc.sections[0].header.part._element
        if header is True
        else doc._part._element
    )
    # TODO : Voir pour chercher également bookmarkend. s'en servir pour supprimer le contenue entre start et end
    bookmarks_list = doc_element.findall(".//" + qn("w:bookmarkStart"))

    for bookmark in bookmarks_list:
        name = bookmark.get(qn("w:name"))
        if name == bookmark_name:
            par = bookmark.getparent()
            bk_id = bookmark.get(qn("w:id"))
            bookmark_end_list = par.findall(".//" + qn("w:bookmarkEnd"))
            print(bookmark_end_list)
            for bk_end in bookmark_end_list:
                if bk_end.get(qn("w:id")) == bk_id:
                    print("on a trouvé le end du bon bk")
                    pos_end = par.index(bk_end)
                    print(pos_end)

            if not isinstance(par, CT_P):
                return False
            else:
                # for elem in par.iter():
                #     print("%s - %s" % (elem.tag, elem.text))
                #print(et.tostring(par, pretty_print=True))
                # Todo : Voir si on ne peut pas plutot modifier le text du bookmark
                #print(par.index(bookmark))
                # for z in range(par.index(bookmark),pos_end):
                #     print(et.tostring(par[z]))
                #     if isinstance(par[z], CT_R):
                #         print('cest un run')

                i = par.index(bookmark) - 1
                p = doc.add_paragraph()
                run = p.add_run(text, style)
                run.underline = underline
                run.italic = italic
                run.bold = bold
                run.font.size = font_size
                run.font.name = font_name
                par.insert(i, run._element)
                p = p._element
                p.getparent().remove(p)
                p._p = p._element = None
                # Essai pour suppresion bookmarks + text
                for i in range(par.index(bookmark), pos_end +1):
                    print(i)
                    asup = par[i]

                    print(et.tostring(asup, pretty_print=True))
                    par.remove(asup)

                # par.remove(bookmark)
                # print(et.tostring(par, pretty_print=True))
                return True
    return False


def loop_bookmark_text(
        doc,
        bookmark_name,
        text,
        underline=False,
        italic=False,
        bold=False,
        style=None,
        header=False,
        font_name=None,
        font_size=None):
    doc_element = (
        doc.sections[0].header.part._element
        if header is True
        else doc._part._element
    )
    # TODO : Voir pour chercher également bookmarkend. s'en servir pour supprimer le contenue entre start et end
    bookmarks_list = doc_element.findall(".//" + qn("w:bookmarkStart"))
    print(bookmarks_list)
    for bookmark in bookmarks_list:
        name = bookmark.get(qn("w:name"))
        print(name)
        par = bookmark.getparent()
        bk_id = bookmark.get(qn("w:id"))
        bookmark_end_list = par.findall(".//" + qn("w:bookmarkEnd"))
        print(bookmark_end_list)
        for bk_end in bookmark_end_list:
            if bk_end.get(qn("w:id")) == bk_id:
                print("on a trouvé le end du bon bk")
                pos_end = par.index(bk_end)
                print(par.index(bookmark))
                print(pos_end)






doc = Document(".\ExtractionLims\Test Rapport Essai corrosion SSC - A22_00202 - 2022-07-05.docx")

#test = loop_bookmark_text(doc, "BK_Condition_Solution", " 1111")

test = bookmark_text(doc, "BK_Condition_Solution", " 1111", font_name='Arial', font_size=127000)

doc.save('testbk.docx')