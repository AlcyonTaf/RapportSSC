import docx
from lxml import etree as et
import pandas as pd
from docx.document import Document

try:
    document = Document()
except TypeError:
    from docx import Document

    document = Document()

from docx.oxml.shared import qn
from docx.oxml.text.paragraph import CT_P
data =[["R0C0","R0C1", "R0C2"],["R1C0","R1C1", "R1C2"],["R2C0","R2C1", "R2C2"]]

print(data)
df = pd.DataFrame(data)

print(df)

doc = Document()
# Apparement on peu merge 2 cell cote a cote de cette facon, mais ensuite cela doit poser probleme pour y acceder
table = doc.add_table(3,3)

for rowindex, row in df.iterrows():
    for colindex, value in row.items():
        table.cell(rowindex,colindex).text = value


# table.cell(0, 0).text ="R0C0"
# #table.cell(0, 1).text ="R0C1"
# table.cell(1, 0).text ="R1C0"
# table.cell(1, 1).text ="R1C1"
table.style = 'Normal Table'

# merge par ligne
# a = table.cell(0,0)
# b = table.cell(0,1)
# A = a.merge(b)

# merge par colonne
a = table.cell(0,0)
b = table.cell(2,0)
A = a.merge(b)



doc.save('testmegercell.docx')
